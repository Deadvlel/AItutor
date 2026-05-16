import io
import json
from sqlalchemy.orm import Session
from models.course import chuDe, taiLieu, loaiCauHoi, cauHoi
from services.ai_service import hoi_gia_su
ID_DO_BAI = 2


def ensure_loai_cau_hoi(db: Session):
    for id_, ten, mo_ta in [
        (1, "kiem_tra", "Câu hỏi trắc nghiệm"),
        (2, "do_bai",   "Câu hỏi dò bài — giọng nói"),
    ]:
        if not db.query(loaiCauHoi).filter(loaiCauHoi.id_loaiCauHoi == id_).first():
            db.add(loaiCauHoi(id_loaiCauHoi=id_, tenLoai=ten, moTa=mo_ta))
    db.commit()


def get_or_create_chu_de(db: Session, ten: str) -> int:
    cd = db.query(chuDe).filter(chuDe.ten_chuDe == ten.strip()).first()
    if not cd:
        cd = chuDe(ten_chuDe=ten.strip())
        db.add(cd)
        db.flush()
    return cd.id_chuDe


def _detect_loai(ten: str) -> str:
    ten = ten.lower()
    if "toán" in ten or "toan" in ten:  return "toan"
    if "văn"  in ten or "van"  in ten:  return "ngu_van"
    if "lý"   in ten or "ly"   in ten:  return "vat_ly"
    if "hóa"  in ten or "hoa"  in ten:  return "hoa_hoc"
    if "sinh" in ten:                    return "sinh_hoc"
    if "sử"   in ten or "su"   in ten:  return "lich_su"
    return "khac"


def _ai_sinh_cau_hoi(noi_dung: str, so_cau: int) -> list[dict]:
    MAX_CHARS = 6000
    if len(noi_dung) > MAX_CHARS:
        noi_dung = noi_dung[:4000] + "\n...\n" + noi_dung[-2000:]

    prompt = f"""Dựa vào nội dung bài học sau, hãy tạo {so_cau} câu hỏi dò bài.

NỘI DUNG:
{noi_dung}

Trả về JSON thuần (không markdown, không giải thích thêm):
{{
  "cau_hoi": [
    {{
      "noi_dung":   "Câu hỏi ngắn gọn rõ ràng",
      "dap_an_mau": "Đáp án đầy đủ để AI so sánh",
      "goi_y":      "Gợi ý ngắn nếu học sinh trả lời sai",
      "thu_tu":     1
    }}
  ]
}}

Yêu cầu: tiếng Việt, bám sát nội dung, đáp án rõ ràng."""

    raw = hoi_gia_su(prompt)
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]

    data = json.loads(raw)
    return data.get("cau_hoi", [])


def _luu_vao_db(
    db:       Session,
    id_cd:    int,
    tieu_de:  str,
    loai:     str,
    do_kho:   int,
    cac_cau:  list[dict],
    ten_file: str = None,
) -> dict:
    tl = db.query(taiLieu).filter(
        taiLieu.tieuDe   == tieu_de,
        taiLieu.id_chuDe == id_cd,
    ).first()
    if not tl:
        tl = taiLieu(
            id_chuDe=id_cd,
            tieuDe=tieu_de,
            loai=loai,
            file=ten_file,
            mucDoKho=do_kho,
        )
        db.add(tl)
        db.flush()

    for ch in cac_cau:
        db.add(cauHoi(
            id_taiLieu    = tl.id_taiLieu,
            id_chuDe      = id_cd,
            id_loaiCauHoi = ID_DO_BAI,
            id_baiKiemTra = None,
            noiDung       = ch.get("noi_dung") or ch.get("cau_hoi", ""),
            dapAnMau      = ch.get("dap_an_mau", ""),
            loiGiaiThich  = ch.get("dap_an_mau", ""),
            goiY          = ch.get("goi_y") or None,
            thuTu         = ch.get("thu_tu", 1),
        ))

    db.commit()
    return {
        "id_taiLieu":  tl.id_taiLieu,
        "tieu_de":     tieu_de,
        "da_them":     len(cac_cau),
        "cac_cau_hoi": [c.get("noi_dung", c.get("cau_hoi", "")) for c in cac_cau],
    }


def xu_ly_excel(file_bytes: bytes, ten_file: str, db: Session) -> dict:
    """
    Cột bắt buộc: chu_de | tieu_de | cau_hoi | dap_an_mau
    Cột tuỳ chọn: goi_y | do_kho | thu_tu
    """
    try:
        import pandas as pd
    except ImportError:
        raise RuntimeError("pip install pandas openpyxl xlrd")

    ensure_loai_cau_hoi(db)

    ten = ten_file.lower()

    if ten.endswith(".csv"):
        df = pd.read_csv(io.BytesIO(file_bytes), encoding="utf-8-sig")
    elif ten.endswith(".xls"):
        df = pd.read_excel(io.BytesIO(file_bytes), engine="xlrd")
    else:
        df = pd.read_excel(io.BytesIO(file_bytes))

    thieu = [c for c in ["chu_de", "tieu_de", "cau_hoi", "dap_an_mau"] if c not in df.columns]
    if thieu:
        raise ValueError(f"File thiếu cột: {', '.join(thieu)}")

    da_them   = 0
    tai_lieus = {}

    for _, row in df.iterrows():
        ten_cd   = str(row["chu_de"]).strip()
        tieu_de  = str(row["tieu_de"]).strip()
        noi_dung = str(row["cau_hoi"]).strip()
        dap_an   = str(row["dap_an_mau"]).strip()
        goi_y    = str(row.get("goi_y", "")).strip()
        do_kho   = int(row["do_kho"]) if "do_kho" in row and str(row["do_kho"]) != "nan" else 1
        thu_tu   = int(row["thu_tu"]) if "thu_tu" in row and str(row["thu_tu"]) != "nan" else da_them + 1
        if not noi_dung or noi_dung.lower() == "nan":
            continue

        id_cd = get_or_create_chu_de(db, ten_cd)
        key   = (ten_cd, tieu_de)

        if key not in tai_lieus:
            tl = db.query(taiLieu).filter(
                taiLieu.tieuDe == tieu_de, taiLieu.id_chuDe == id_cd
            ).first()
            if not tl:
                tl = taiLieu(
                    id_chuDe=id_cd, tieuDe=tieu_de,
                    loai=_detect_loai(ten_cd), mucDoKho=do_kho,
                )
                db.add(tl)
                db.flush()
            tai_lieus[key] = tl.id_taiLieu

        db.add(cauHoi(
            id_taiLieu    = tai_lieus[key],
            id_chuDe      = id_cd,
            id_loaiCauHoi = ID_DO_BAI,
            id_baiKiemTra = None,
            noiDung       = noi_dung,
            dapAnMau      = dap_an,
            loiGiaiThich  = dap_an,
            goiY          = goi_y if goi_y and goi_y.lower() != "nan" else None,
            thuTu         = thu_tu,
        ))
        da_them += 1

    db.commit()
    return {
        "da_them":  da_them,
        "bai_hocs": list({k[1] for k in tai_lieus}),
    }


def xu_ly_pdf(
    file_bytes: bytes,
    ten_chu_de: str,
    tieu_de:    str,
    so_cau:     int,
    db:         Session,
) -> dict:
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pip install pdfplumber")

    ensure_loai_cau_hoi(db)

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        noi_dung = "\n".join(p.extract_text() or "" for p in pdf.pages)

    if len(noi_dung.strip()) < 50:
        raise ValueError(
            "PDF không đọc được text — có thể là PDF scan (hình ảnh). "
            "Hãy dùng PDF có chứa text thật."
        )

    cac_cau = _ai_sinh_cau_hoi(noi_dung, so_cau)
    if not cac_cau:
        raise ValueError("AI không sinh được câu hỏi. Thử giảm số câu hoặc kiểm tra nội dung PDF.")

    id_cd = get_or_create_chu_de(db, ten_chu_de)
    return _luu_vao_db(db, id_cd, tieu_de, _detect_loai(ten_chu_de), 2, cac_cau)


def xu_ly_word(
    file_bytes: bytes,
    ten_chu_de: str,
    tieu_de:    str,
    so_cau:     int,
    db:         Session,
) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("pip install python-docx")

    ensure_loai_cau_hoi(db)

    doc      = Document(io.BytesIO(file_bytes))
    noi_dung = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            noi_dung += "\n" + " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )

    if len(noi_dung.strip()) < 50:
        raise ValueError("File Word trống hoặc không có nội dung text.")

    cac_cau = _ai_sinh_cau_hoi(noi_dung, so_cau)
    if not cac_cau:
        raise ValueError("AI không sinh được câu hỏi. Thử giảm số câu hoặc kiểm tra nội dung.")

    id_cd = get_or_create_chu_de(db, ten_chu_de)
    return _luu_vao_db(db, id_cd, tieu_de, _detect_loai(ten_chu_de), 2, cac_cau)